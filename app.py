from flask import Flask, render_template, request, redirect, url_for, send_file
from datetime import datetime
import io
import csv
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle
from openpyxl import Workbook
from docx import Document
from docx.shared import Inches

app = Flask(__name__)

# In-memory storage for teams (replace with a database in a real application)
teams = []

def calculate_total_points(position_points, finish_points):
    return position_points + finish_points

def sort_teams(teams):
    return sorted(teams, key=lambda x: x['total_points'], reverse=True)

@app.route('/')
def index():
    sorted_teams = sort_teams(teams)
    return render_template('index.html', teams=sorted_teams)

@app.route('/add_team', methods=['POST'])
def add_team():
    name = request.form['name']
    matches_played = int(request.form['matches_played'])
    position_points = int(request.form['position_points'])
    finish_points = int(request.form['finish_points'])
    chicken_dinners = int(request.form['chicken_dinners'])
    
    total_points = calculate_total_points(position_points, finish_points)
    
    team = {
        'id': datetime.now().timestamp(),
        'name': name,
        'matches_played': matches_played,
        'position_points': position_points,
        'finish_points': finish_points,
        'total_points': total_points,
        'chicken_dinners': chicken_dinners
    }
    
    teams.append(team)
    return redirect(url_for('index'))

@app.route('/export/<format>')
def export(format):
    sorted_teams = sort_teams(teams)
    if format == 'pdf':
        return export_pdf(sorted_teams)
    elif format == 'excel':
        return export_excel(sorted_teams)
    elif format == 'word':
        return export_word(sorted_teams)
    else:
        return "Invalid format", 400

def export_pdf(teams):
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    elements = []

    data = [['Rank', 'Team Name', 'Matches', 'Position Points', 'Finish Points', 'Total Points', 'Chicken Dinners']]
    for i, team in enumerate(teams, 1):
        data.append([
            i,
            team['name'],
            team['matches_played'],
            team['position_points'],
            team['finish_points'],
            team['total_points'],
            team['chicken_dinners']
        ])

    table = Table(data)
    table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 14),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
        ('TEXTCOLOR', (0, 1), (-1, -1), colors.black),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 1), (-1, -1), 12),
        ('TOPPADDING', (0, 1), (-1, -1), 6),
        ('BOTTOMPADDING', (0, 1), (-1, -1), 6),
        ('GRID', (0, 0), (-1, -1), 1, colors.black)
    ]))

    elements.append(table)
    doc.build(elements)

    buffer.seek(0)
    return send_file(buffer, as_attachment=True, download_name='bgmi_standings.pdf', mimetype='application/pdf')

def export_excel(teams):
    wb = Workbook()
    ws = wb.active
    ws.title = "BGMI Standings"

    headers = ['Rank', 'Team Name', 'Matches', 'Position Points', 'Finish Points', 'Total Points', 'Chicken Dinners']
    ws.append(headers)

    for i, team in enumerate(teams, 1):
        ws.append([
            i,
            team['name'],
            team['matches_played'],
            team['position_points'],
            team['finish_points'],
            team['total_points'],
            team['chicken_dinners']
        ])

    buffer = io.BytesIO()
    wb.save(buffer)
    buffer.seek(0)

    return send_file(buffer, as_attachment=True, download_name='bgmi_standings.xlsx', mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')

def export_word(teams):
    doc = Document()
    doc.add_heading('BGMI Standings', 0)

    table = doc.add_table(rows=1, cols=7)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Rank'
    hdr_cells[1].text = 'Team Name'
    hdr_cells[2].text = 'Matches'
    hdr_cells[3].text = 'Position Points'
    hdr_cells[4].text = 'Finish Points'
    hdr_cells[5].text = 'Total Points'
    hdr_cells[6].text = 'Chicken Dinners'

    for i, team in enumerate(teams, 1):
        row_cells = table.add_row().cells
        row_cells[0].text = str(i)
        row_cells[1].text = team['name']
        row_cells[2].text = str(team['matches_played'])
        row_cells[3].text = str(team['position_points'])
        row_cells[4].text = str(team['finish_points'])
        row_cells[5].text = str(team['total_points'])
        row_cells[6].text = str(team['chicken_dinners'])

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return send_file(buffer, as_attachment=True, download_name='bgmi_standings.docx', mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

if __name__ == '__main__':
    app.run(debug=True)
